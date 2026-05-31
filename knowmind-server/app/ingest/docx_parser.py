from __future__ import annotations

import logging
import os
import re
import tempfile
import xml.etree.ElementTree as ET
import zipfile
from pathlib import Path

from app.ingest.types import PageText, ParseResult

log = logging.getLogger(__name__)

_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_W_TAG = f"{{{_W_NS}}}"
_NULL_REL_PATTERN = re.compile(
    r'<Relationship\b[^>]*\bTarget\s*=\s*"(?:NULL|null|)"[^>]*/>\s*',
    re.IGNORECASE,
)
_NULL_REL_PATTERN_SQ = re.compile(
    r"<Relationship\b[^>]*\bTarget\s*=\s*'(?:NULL|null|)'[^>]*/>\s*",
    re.IGNORECASE,
)


def _cell_text(cell) -> str:
    return (cell.text or "").strip()


def _table_to_markdown(table) -> str:
    rows: list[list[str]] = []
    for row in table.rows:
        cells = [_cell_text(c) for c in row.cells]
        if any(cells):
            rows.append(cells)
    if not rows:
        return ""
    col_count = max(len(r) for r in rows)
    normalized = [r + [""] * (col_count - len(r)) for r in rows]
    lines: list[str] = []
    for i, row in enumerate(normalized):
        line = "| " + " | ".join(c.replace("|", "\\|") for c in row) + " |"
        lines.append(line)
        if i == 0:
            lines.append("| " + " | ".join("---" for _ in row) + " |")
    return "\n".join(lines)


def _heading_level(style_name: str | None) -> int | None:
    if not style_name:
        return None
    name = style_name.lower()
    if name.startswith("heading"):
        try:
            return int(name.replace("heading", "").strip() or "1")
        except ValueError:
            return 1
    if "标题" in style_name:
        for n in range(1, 7):
            if str(n) in style_name:
                return n
        return 1
    return None


def strip_null_relationships(rels_xml: bytes) -> bytes:
    """移除 OOXML 中 Target=NULL 的损坏 Relationship（部分 WPS/旧版 Word 导出会带）。"""
    text = rels_xml.decode("utf-8", errors="ignore")
    text = _NULL_REL_PATTERN.sub("", text)
    text = _NULL_REL_PATTERN_SQ.sub("", text)
    return text.encode("utf-8")


def repair_docx_file(path: str) -> str:
    """复制 docx 并清理 .rels 中的 NULL 引用，返回临时文件路径（调用方负责删除）。"""
    fd, out_path = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    with zipfile.ZipFile(path, "r") as zin, zipfile.ZipFile(out_path, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename.endswith(".rels") and b"NULL" in data:
                data = strip_null_relationships(data)
            zout.writestr(item, data)
    return out_path


def _parse_docx_xml_fallback(path: str) -> str:
    """不依赖 python-docx，直接从 word/document.xml 抽取段落文本。"""
    with zipfile.ZipFile(path, "r") as zf:
        try:
            xml_bytes = zf.read("word/document.xml")
        except KeyError as e:
            raise ValueError("DOCX 缺少 word/document.xml，文件可能已损坏") from e
    root = ET.fromstring(xml_bytes)
    blocks: list[str] = []
    for para in root.iter(f"{_W_TAG}p"):
        parts: list[str] = []
        for node in para.iter(f"{_W_TAG}t"):
            if node.text:
                parts.append(node.text)
        line = "".join(parts).strip()
        if line:
            blocks.append(line)
    return "\n\n".join(blocks)


def _build_parse_result(content: str, path: str, filename: str | None) -> ParseResult:
    base = (filename or Path(path).name).rsplit(".", 1)[0]
    title = base[:200] if base else None
    summary = None
    for block in content.split("\n\n"):
        t = block.lstrip("#").strip()
        if t:
            summary = t[:500]
            break
    pages = [PageText(page_index=0, text=content)] if content else []
    return ParseResult(pages=pages, title=title, summary=summary, content=content)


def _parse_docx_with_python_docx(path: str, filename: str | None = None) -> ParseResult:
    from docx import Document as DocxDocument  # noqa: PLC0415
    from docx.table import Table  # noqa: PLC0415
    from docx.text.paragraph import Paragraph  # noqa: PLC0415

    doc = DocxDocument(path)
    blocks: list[str] = []
    body = doc.element.body
    para_idx = 0
    table_idx = 0

    for child in body.iterchildren():
        tag = child.tag.split("}")[-1]
        if tag == "p":
            if para_idx >= len(doc.paragraphs):
                continue
            para: Paragraph = doc.paragraphs[para_idx]
            para_idx += 1
            text = (para.text or "").strip()
            if not text:
                continue
            level = _heading_level(para.style.name if para.style else None)
            if level:
                blocks.append(f"{'#' * level} {text}")
            else:
                blocks.append(text)
        elif tag == "tbl":
            if table_idx >= len(doc.tables):
                continue
            table: Table = doc.tables[table_idx]
            table_idx += 1
            md = _table_to_markdown(table)
            if md:
                blocks.append(md)

    content = "\n\n".join(blocks)
    return _build_parse_result(content, path, filename)


def parse_docx(path: str, filename: str | None = None) -> ParseResult:
    last_err: Exception | None = None

    try:
        return _parse_docx_with_python_docx(path, filename)
    except Exception as e:
        last_err = e
        log.warning("docx python-docx failed (%s): %s", e, path)

    repaired: str | None = None
    try:
        repaired = repair_docx_file(path)
        try:
            return _parse_docx_with_python_docx(repaired, filename)
        except Exception as e:
            last_err = e
            log.warning("docx repair+python-docx failed (%s): %s", e, path)
        for src in (repaired, path):
            try:
                content = _parse_docx_xml_fallback(src)
            except Exception as e:
                last_err = e
                continue
            if content.strip():
                log.info("docx parsed via xml fallback: %s", path)
                return _build_parse_result(content, path, filename)
    finally:
        if repaired and os.path.isfile(repaired):
            try:
                os.unlink(repaired)
            except OSError:
                pass

    msg = "无法解析该 DOCX（文件可能含损坏引用）。请用 Word/WPS「另存为」新 .docx 后重试"
    raise ValueError(msg) from last_err


def parse_doc(path: str, filename: str | None = None) -> ParseResult:
    """旧版 .doc：尝试 olefile 提取，失败则提示转换。"""
    try:
        import olefile  # noqa: PLC0415

        if not olefile.isOleFile(path):
            raise ValueError("不是有效的 Word .doc 文件")
        ole = olefile.OleFileIO(path)
        try:
            if ole.exists("WordDocument"):
                data = ole.openstream("WordDocument").read()
                text_parts: list[str] = []
                chunk = data.decode("utf-16-le", errors="ignore")
                for line in chunk.split("\x00"):
                    t = line.strip()
                    if len(t) >= 2 and any(c.isalnum() or "\u4e00" <= c <= "\u9fff" for c in t):
                        text_parts.append(t)
                content = "\n\n".join(dict.fromkeys(text_parts))[:50000]
                if len(content) >= 20:
                    base = (filename or Path(path).name).rsplit(".", 1)[0]
                    return ParseResult(
                        pages=[PageText(page_index=0, text=content)],
                        title=base[:200],
                        summary=content[:500],
                        content=content,
                    )
        finally:
            ole.close()
    except ImportError:
        pass
    except Exception:
        pass
    raise ValueError("无法解析 .doc 文件，请另存为 .docx 后重新上传")
